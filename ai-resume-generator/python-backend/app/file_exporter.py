"""
File export utilities for generating PDF and DOCX resume files.
Produces properly formatted, professional-looking documents.
"""

import logging
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from .config import settings

logger = logging.getLogger(__name__)


def _ensure_output_dir():
    """Create the output directory if it doesn't exist."""
    os.makedirs(settings.output_dir, exist_ok=True)


def export_pdf(content: str) -> str:
    """Export resume content to a styled PDF file.

    Returns the path to the generated file.
    """
    _ensure_output_dir()
    output_path = os.path.join(settings.output_dir, "resume.pdf")

    try:
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            topMargin=0.6 * inch,
            bottomMargin=0.6 * inch,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        style_name = ParagraphStyle(
            "ResumeName",
            parent=styles["Title"],
            fontSize=18,
            leading=22,
            spaceAfter=4,
            textColor=HexColor("#1a1a2e"),
            alignment=1,  # CENTER
        )
        style_contact = ParagraphStyle(
            "ResumeContact",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=12,
            textColor=HexColor("#555555"),
            alignment=1,
        )
        style_heading = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading2"],
            fontSize=12,
            leading=16,
            spaceBefore=14,
            spaceAfter=6,
            textColor=HexColor("#1e40af"),
            borderWidth=0,
            borderPadding=0,
        )
        style_body = ParagraphStyle(
            "ResumeBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=4,
            textColor=HexColor("#333333"),
        )

        flowables = []
        lines = content.split("\n")

        # Known section headings
        section_keywords = {
            "CAREER OBJECTIVE", "OBJECTIVE", "SKILLS", "EXPERIENCE",
            "EDUCATION", "PROJECTS", "LINKS", "SUMMARY", "CERTIFICATIONS",
        }

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                flowables.append(Spacer(1, 6))
                continue

            if i == 0:
                # First line = name
                flowables.append(Paragraph(stripped, style_name))
            elif i == 1:
                # Second line = contact info
                flowables.append(Paragraph(stripped, style_contact))
            elif i == 2 and "|" in stripped or "http" in stripped.lower():
                # Links line
                flowables.append(Paragraph(stripped, style_contact))
            elif stripped.upper() in section_keywords:
                flowables.append(Paragraph(stripped.upper(), style_heading))
            else:
                flowables.append(Paragraph(stripped, style_body))

        doc.build(flowables)
        logger.info("PDF exported to %s", output_path)
        return output_path

    except Exception as e:
        logger.error("PDF export failed: %s", e)
        raise


def export_docx(content: str) -> str:
    """Export resume content to a styled DOCX file.

    Returns the path to the generated file.
    """
    _ensure_output_dir()
    output_path = os.path.join(settings.output_dir, "resume.docx")

    try:
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        lines = content.split("\n")

        section_keywords = {
            "CAREER OBJECTIVE", "OBJECTIVE", "SKILLS", "EXPERIENCE",
            "EDUCATION", "PROJECTS", "LINKS", "SUMMARY", "CERTIFICATIONS",
        }

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue

            if i == 0:
                # Name heading
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            elif i == 1:
                # Contact info
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif i == 2 and ("|" in stripped or "http" in stripped.lower()):
                # Links
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif stripped.upper() in section_keywords:
                # Section heading
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(stripped.upper())
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            else:
                # Body text
                p = doc.add_paragraph(stripped)

        doc.save(output_path)
        logger.info("DOCX exported to %s", output_path)
        return output_path

    except Exception as e:
        logger.error("DOCX export failed: %s", e)
        raise
